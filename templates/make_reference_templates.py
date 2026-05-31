"""Generates the three *reference* templates with placeholder markup.

These are NOT meant to ship as-is — they show you exactly what Jinja/docxtpl
markup to paste into your real company-branded Technical Summary / Code
Comparison / Change Request files. Run once; then mirror the placeholders into
your branded copies.

Layout (Word docs):
  Page 1  — Title + automatic Table of Contents (Word field, updates on open)
  Page 2+ — content with real Heading styles so the TOC populates
"""
import os

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

HERE = os.path.dirname(os.path.abspath(__file__))


def _toc_field(doc, heading_levels="1-3"):
    """Insert a real Word Table of Contents field (Automatic Table).
    Word builds/refreshes it from the Heading styles when the doc is opened
    and fields are updated."""
    p = doc.add_paragraph()
    run = p.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\o "{heading_levels}" \\h \\z \\u '
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    run._r.append(begin); run._r.append(instr); run._r.append(sep)
    hint = p.add_run("Right-click here and choose “Update Field” to build the table of contents.")
    hint.italic = True
    end_run = p.add_run()
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _force_update_fields(doc):
    """Make Word offer to update fields (the TOC) when the document is opened."""
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


def _title_page(doc, title_text):
    doc.add_paragraph(title_text, style="Title")
    # NOT a Heading style — otherwise this label would list itself inside the TOC.
    label = doc.add_paragraph()
    run = label.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(14)
    _toc_field(doc, "1-3")
    doc.add_page_break()


def make_technical_summary():
    doc = Document()
    _title_page(doc, "Technical Summary - {{ ticket_ids }}")

    # One section per ticket; each ticket starts on a new page.
    doc.add_paragraph("{%p for t in tickets %}")
    doc.add_paragraph("{%p if not loop.first %}")
    doc.add_page_break()
    doc.add_paragraph("{%p endif %}")
    doc.add_paragraph("{{ t.id }}: {{ t.title }}", style="Heading 2")
    for n, (q, field) in enumerate([
        ("What is the problem?", "t.problem"),
        ("How does this affect the users?", "t.user_impact"),
        ("Solution Implemented", "t.solution"),
        ("How was this tested?", "t.testing"),
    ], start=1):
        doc.add_paragraph(f"{n}) {q}", style="Heading 3")
        doc.add_paragraph("{{ %s }}" % field)
    doc.add_paragraph("{%p endfor %}")

    _force_update_fields(doc)
    doc.save(os.path.join(HERE, "technical_summary_template.docx"))


def make_code_comparison():
    doc = Document()
    _title_page(doc, "Code Comparison - {{ ticket_ids }}")

    # Page 2: introduction + the list of tickets covered.
    doc.add_paragraph("INTRODUCTION", style="Heading 2")
    doc.add_paragraph(
        "This document will present the Code Comparison of the changes made in "
        "the following items:"
    )
    doc.add_paragraph("{%p for t in tickets %}")
    doc.add_paragraph("{{ t.id }} - {{ t.title }}", style="List Bullet")
    doc.add_paragraph("{%p endfor %}")
    doc.add_page_break()

    # Page 3+: per-repo comparison (each repo on its own page).
    doc.add_paragraph("{%p for repo in repos %}")
    doc.add_paragraph("{%p if not loop.first %}")
    doc.add_page_break()
    doc.add_paragraph("{%p endif %}")
    doc.add_paragraph("{{ repo.name }}", style="Heading 2")
    doc.add_paragraph("Version to be deployed: {{r repo.deploy_link }}", style="List Bullet")
    doc.add_paragraph("Current PROD version: {{r repo.prod_link }}", style="List Bullet")
    doc.add_paragraph("Comparison Link: {{r repo.compare_link }}", style="List Bullet")
    doc.add_paragraph("Files changed", style="Heading 3")
    doc.add_paragraph("{{ repo.stat }}")
    doc.add_paragraph("Commits", style="Heading 3")
    doc.add_paragraph("{{r repo.commits_rich }}")
    doc.add_paragraph("Code changes", style="Heading 3")
    doc.add_paragraph("{{r repo.diff_rich }}")
    doc.add_paragraph("{%p endfor %}")

    _force_update_fields(doc)
    doc.save(os.path.join(HERE, "code_comparison_template.docx"))


def make_change_request():
    wb = Workbook()
    ws = wb.active
    ws.title = "Change Request"
    ws.column_dimensions["A"].width = 42
    ws.column_dimensions["B"].width = 80

    header_fill = PatternFill("solid", fgColor="D9E1F2")
    ws["A1"] = "Field"
    ws["B1"] = "Value"
    for c in ("A1", "B1"):
        ws[c].font = Font(bold=True)
        ws[c].fill = header_fill

    rows = [
        ("Change Request Number", "{{ cr_number }}"),
        ("Deployment Date", "{{ deploy_date }}"),
        ("Tickets", "{{ ticket_ids }}"),
        ("Repositories Touched", "{{ repo_names }}"),
        ("Versions to be Deployed", "{{ deploy_versions }}"),
        ("Backout Versions", "{{ backout_versions }}"),
        ("Business Justification", "{{ business_justification }}"),
        ("Summary of Change", "{{ change_summary }}"),
        ("What happens if the change is NOT deployed?", "{{ impact_if_not_deployed }}"),
        ("Worst case if deployed and something goes wrong?", "{{ worst_case }}"),
        ("Rollback / Backout Plan", "{{ rollback_plan }}"),
        ("Deployment Steps", "{{ deployment_steps | lines }}"),
    ]
    for i, (q, a) in enumerate(rows, start=2):
        ws.cell(row=i, column=1, value=q).font = Font(bold=True)
        cell = ws.cell(row=i, column=2, value=a)
        cell.alignment = Alignment(wrap_text=True, vertical="top")

    wb.save(os.path.join(HERE, "change_request_template.xlsx"))


if __name__ == "__main__":
    make_technical_summary()
    make_code_comparison()
    make_change_request()
    print("Reference templates written to", HERE)
